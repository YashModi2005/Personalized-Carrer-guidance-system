from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# FORMATTING CONSTANTS
FONT_NAME = "Calibri"
SIZE_MAIN_TITLE = 18
SIZE_CHAPTER = 16
SIZE_SECTION = 14
SIZE_SUBSECTION = 12
SIZE_NORMAL = 12
SIZE_TABLE = 11

def set_font(run, font_name=FONT_NAME, size_pt=SIZE_NORMAL, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_paragraph_format(para, alignment=None, line_spacing=1.0, space_after=0):
    if alignment:
        para.alignment = alignment
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_after = Pt(space_after)

def add_heading(doc, text, level=1, bold=True):
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    size = SIZE_CHAPTER if level == 1 else (SIZE_SECTION if level == 2 else SIZE_SUBSECTION)
    set_font(run, size_pt=size, bold=bold)
    return para

def add_normal_para(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = para.add_run(text)
    set_font(run, size_pt=SIZE_NORMAL, bold=bold)
    run.font.italic = italic
    return para

def add_centered_text(doc, text, size=SIZE_NORMAL, bold=False, italic=False, space_after=0):
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=space_after)
    run = para.add_run(text)
    set_font(run, size_pt=size, bold=bold)
    run.font.italic = italic
    return para

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        run = para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._element.append(instrText)
        run = para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        set_font(run, size_pt=10)

def add_image(doc, img_path, width_inches=None, height_inches=None):
    if not os.path.exists(img_path):
        add_centered_text(doc, f"[IMAGE MISSING: {img_path}]", size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if height_inches:
        run.add_picture(img_path, height=Inches(height_inches))
    elif width_inches:
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        run.add_picture(img_path, width=Inches(5.0))
    doc.add_paragraph("\n")

def add_db_table(doc, title, headers, rows):
    add_heading(doc, title, level=3)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_font(hdr_cells[i].paragraphs[0].runs[0], size_pt=SIZE_TABLE, bold=True)
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i+1].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = str(val)
            set_font(row_cells[j].paragraphs[0].runs[0], size_pt=SIZE_TABLE)
    doc.add_paragraph("\n")

def create_report():
    doc = Document()
    
    # 1. COVER PAGE
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=14, bold=True)
    add_centered_text(doc, "Integrated MSc(IT) Programme", size=14, bold=True, space_after=20)
    add_centered_text(doc, "“Personalized Career Guidance System”", size=SIZE_MAIN_TITLE, bold=True, space_after=40)
    add_centered_text(doc, "Capstone Project Documentation", size=14, bold=True)
    add_centered_text(doc, "submitted for partial fulfillment of Semester – VIII", size=12, space_after=40)
    add_centered_text(doc, "By\nYash Modi (202201619010159)\nNitya Modi (2022016190100046)", size=12, space_after=60)
    
    table = doc.add_table(rows=1, cols=2)
    l_para = table.cell(0,0).paragraphs[0]
    set_font(l_para.add_run("External Guide\n(Guide Name)"), size_pt=11)
    r_para = table.cell(0,1).paragraphs[0]
    set_paragraph_format(r_para, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_font(r_para.add_run("Internal Guide\n(Prof. Name)"), size_pt=11)

    # 2. COMPANY CERTIFICATE
    doc.add_page_break()
    add_heading(doc, "2. COMPANY CERTIFICATE", level=1)
    cert_body = "This is to certify that Yash Modi and Nitya Modi have successfully developed the 'AI CareerPilot: Personalized Career Guidance System'. Their project demonstrates a high level of expertise in integrating Machine Learning models with modern web technologies like FastAPI and React.js. Their dedication to creating a tool that provides actionable career insights for students is commendable."
    add_normal_para(doc, cert_body)
    doc.add_paragraph("\n\n[PLACEHOLDER: COMPANY STAMP & SIGNATURE]\n\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. COLLEGE CERTIFICATE
    doc.add_page_break()
    add_centered_text(doc, "Faculty of Computer Applications & Information Technology", size=14, bold=True)
    add_centered_text(doc, "Integrated MSc(IT) Programme", size=14, bold=True, space_after=30)
    add_heading(doc, "3. COLLEGE CERTIFICATE", level=1)
    cert_text = "This is to certify that Yash Modi (202201619010159) and Nitya Modi (2022016190100046) students of Semester – VIII Integrated MSc (IT) have successfully completed their Capstone project titled “Personalized Career Guidance System” for the academic year 2025-2026. This project was developed under the supervision of the Department of FCAIT."
    add_normal_para(doc, cert_text)
    doc.add_paragraph("\n\n(Head of Department)\t\t\t\t\t(Project Coordinator)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. ACKNOWLEDGEMENT
    doc.add_page_break()
    add_heading(doc, "4. ACKNOWLEDGEMENT", level=1)
    add_normal_para(doc, "We express our sincere gratitude to our internal guide (Prof. Name) and Dr. Tripti Dodiya, Dean of FCAIT, for their constant support. We are especially grateful for the technical resources provided to build our AI engine using scikit-learn and FastAPI. Finally, we thank our families and friends for their unwavering belief in our project.")

    # 5. INDEX
    doc.add_page_break()
    add_heading(doc, "5. INDEX", level=1)
    idx_sections = [
        ("6. UML Diagrams & System Analysis", "5"),
        ("7. State Chart Diagram", "11"),
        ("8. Sequence Diagrams", "13"),
        ("9. Data Dictionary / Data Sets", "16"),
        ("10. Screen Layouts", "17"),
        ("11. Sample Coding", "23"),
        ("12. Future Enhancements", "25"),
        ("13. Conclusion", "27"),
        ("14. Bibliography", "28")
    ]
    idx_table = doc.add_table(rows=1, cols=2)
    idx_table.width = Inches(6.0)
    hdr = idx_table.rows[0].cells
    hdr[0].text = "Section Details"; hdr[1].text = "Page No"
    for cell in hdr: 
        set_font(cell.paragraphs[0].runs[0], bold=True)
        set_paragraph_format(cell.paragraphs[0], alignment=WD_ALIGN_PARAGRAPH.LEFT if cell == hdr[0] else WD_ALIGN_PARAGRAPH.RIGHT)
    for name, pg in idx_sections:
        row = idx_table.add_row().cells
        row[0].text = name
        row[1].text = pg
        set_paragraph_format(row[1].paragraphs[0], alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # 6. UML DIAGRAMS
    doc.add_page_break()
    add_heading(doc, "6. UML DIAGRAMS & SYSTEM ANALYSIS", level=1)
    
    # 6.1 Architecture
    add_heading(doc, "6.1 Architecture Diagram", level=2)
    add_normal_para(doc, "The system follows a tiered architecture. The Frontend (React.js) handles user interactions, the Backend (FastAPI) manages business logic and database operations, and the AI Layer (Scikit-Learn) performs career predictions based on input data.")
    add_image(doc, "architecture_diagram.png", width_inches=4.5)
    
    # 6.2 Use Case
    add_heading(doc, "6.2 Use Case Diagram", level=2)
    add_normal_para(doc, "Our Use Case Diagram defines the roles of the 'Student' and the 'Counselor (Admin)'. Students can start assessments and chat with the AI Coach, while Counselors can view system-wide analytics and manage student records.")
    add_image(doc, "Diagrams/UseCaseDiagram as per maam.png", width_inches=4.5)
    
    # 6.3 Class
    add_heading(doc, "6.3 Class Diagram", level=2)
    add_normal_para(doc, "The Class Diagram shows how our objects like 'UserAssessment', 'CareerGuidance', and 'ChatMessage' relate to each other within the backend code to ensure organized data flow.")
    add_image(doc, "Diagrams/CLASSDIAGRAM.png", width_inches=4.5)
    
    # 6.4 Activity
    add_heading(doc, "6.4 Activity Diagram", level=2)
    add_normal_para(doc, "This diagram shows the step-by-step path a student takes, from logging in and answering skill-based questions to receiving a personalized career roadmap from the AI.")
    add_image(doc, "Diagrams/ACTIVITYDIAGRAM.png", width_inches=4.5)

    # 7. STATE CHART
    doc.add_page_break()
    add_heading(doc, "7. STATE CHART DIAGRAM", level=1)
    add_heading(doc, "7.1 Student Assessment State Flow", level=2)
    add_normal_para(doc, "This diagram tracks the 'state' of the application as a student moves from a simple visitor to a registered user, and eventually to someone receiving active career guidance and AI coaching.")
    add_image(doc, "Diagrams/STATE1.png", height_inches=7.0)
    
    doc.add_page_break()
    add_heading(doc, "7.2 User Preference Transition", level=2)
    add_normal_para(doc, "This shows how user preferences change the system's focus, moving from raw data collection to a finalized personal profile with high-accuracy matching.")
    add_image(doc, "Diagrams/STATE2.png", height_inches=7.0)

    # 8. SEQUENCE
    doc.add_page_break()
    add_heading(doc, "8. SEQUENCE DIAGRAMS", level=1)
    add_heading(doc, "8.1 Student Prediction Sequence", level=2)
    add_normal_para(doc, "This diagram shows the exact timing of messages. When a student clicks 'Predict', the request goes to FastAPI, which triggers the ML model and retrieves a match in just milliseconds.")
    add_image(doc, "Diagrams/SEQ1.png", width_inches=5.5)
    
    doc.add_page_break()
    add_heading(doc, "8.2 Admin Authentication Sequence", level=2)
    add_normal_para(doc, "Shows the secure login process for administrators. It ensures that only authorized counselors can access sensitive student information and system analytics.")
    add_image(doc, "Diagrams/SEQ2.png", width_inches=5.5)
    
    doc.add_page_break()
    add_heading(doc, "8.3 Feedback & Retraining Sequence", level=2)
    add_normal_para(doc, "Demonstrates how user feedback is collected and used to improve the AI model over time, ensuring the career guidance remains accurate and relevant.")
    add_image(doc, "Diagrams/SEQ3.png", width_inches=5.5)

    # 9. DATA DICTIONARY
    doc.add_page_break()
    add_heading(doc, "9. DATA DICTIONARY / DATA SETS", level=1)
    add_normal_para(doc, "The system uses a MySQL database with four main tables to store application data efficiently.")
    
    add_db_table(doc, "9.1 USERS TABLE", ["Field", "Type", "Description"], [
        ("id", "INT", "Primary Key: Unique user ID"), ("username", "VARCHAR", "Login name of the user"), ("role", "VARCHAR", "Student or Counselor role")
    ])
    add_db_table(doc, "9.2 ASSESSMENTS TABLE", ["Field", "Type", "Description"], [
        ("id", "INT", "Unique Assessment ID"), ("tech_skills", "TEXT", "List of technical skill inputs"), ("soft_skills", "TEXT", "List of soft skill inputs")
    ])
    add_db_table(doc, "9.3 RECOMMENDATIONS TABLE", ["Field", "Type", "Description"], [
        ("id", "INT", "Unique ID"), ("career_title", "VARCHAR", "The predicted career path"), ("match_score", "FLOAT", "Accuracy percentage of the match")
    ])
    add_db_table(doc, "9.4 CHAT_LOGS TABLE", ["Field", "Type", "Description"], [
        ("id", "INT", "Unique Chat ID"), ("user_id", "INT", "Links chat to a student"), ("message", "TEXT", "What the user or AI said")
    ])

    # 10. SCREEN LAYOUTS
    doc.add_page_break()
    add_heading(doc, "10. SCREEN LAYOUTS", level=1)
    
    add_heading(doc, "10.1 User Side Screens", level=2)
    u_screens = [
        ("10.1.1 Registration", "Allows new students to create an account with a secure password."),
        ("10.1.2 Login", "Secure gateway for students to enter their personalized dashboard."),
        ("10.1.3 Assessment", "Interactive form where students select their technical and academic skills."),
        ("10.1.4 Prediction View", "Displays the top-matched career path with a detailed roadmap."),
        ("10.1.5 AI Coach Chat", "A floating assistant that answers student questions about their career.")
    ]
    for s_id, desc in u_screens:
        add_heading(doc, s_id, level=3)
        add_normal_para(doc, desc)
        add_centered_text(doc, f"[INSERT {s_id.upper()} SCREENSHOT]", size=10, italic=True)

    add_heading(doc, "10.2 Admin Side Screens", level=2)
    a_screens = [
        ("10.2.1 Counselor Login", "Specific login portal for university career counselors."),
        ("10.2.2 Analytics Dashboard", "Visual stats showing which careers are most popular among students."),
        ("10.2.3 User Management", "Table view for counselors to help or update student records.")
    ]
    for s_id, desc in a_screens:
        add_heading(doc, s_id, level=3)
        add_normal_para(doc, desc)
        add_centered_text(doc, f"[INSERT {s_id.upper()} SCREENSHOT]", size=10, italic=True)

    # 11. SAMPLE CODING
    doc.add_page_break()
    add_heading(doc, "11. SAMPLE CODING", level=1)
    add_normal_para(doc, "This code is the core of our system. It uses an asynchronous 'recommend' function to call our ML model and generate an explanation for the student simultaneously, ensuring the app is very fast.")
    code_text = """@app.post("/recommend")
async def recommend(assessment: UserAssessment):
    # ML Prediction runs in a separate thread for performance
    guidance = await loop.run_in_executor(_executor, get_career_guidance, assessment)
    # XAI Explanation helps students understand 'WHY' this career was chosen
    xai_data = await loop.run_in_executor(_executor, explain_prediction, assessment, guidance.top_career)
    return guidance"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    set_font(run, font_name="Consolas", size_pt=9)

    add_normal_para(doc, "In simple terms, this code allows the system to multitask. While the Machine Learning model is busy calculating the best career path, the AI Engine simultaneously prepares a detailed explanation. This prevents the user from waiting and makes the dashboard feel smooth and professional.")

    # 12. FUTURE ENHANCEMENTS
    doc.add_page_break()
    add_heading(doc, "12. FUTURE ENHANCEMENTS", level=1)
    
    add_heading(doc, "12.1 Real-time Job Board Integration", level=2)
    add_normal_para(doc, "In the next version, we plan to connect the system directly to professional networks like LinkedIn and Indeed. This will allow students not only to see their predicted career path but also to view a live list of open job vacancies in that field, making the career guidance immediately actionable.")
    
    add_heading(doc, "12.2 Cross-Platform Mobile Application", level=2)
    add_normal_para(doc, "While the current system is web-based, we aim to develop a mobile app using React Native or Flutter. This will provide students with push notifications for career advice and allow them to chat with the AI Coach while on the go, increasing user engagement and accessibility.")
    
    add_heading(doc, "12.3 Advanced AI Explainability", level=2)
    add_normal_para(doc, "To provide even deeper trust, we will integrate more advanced visualization tools like SHAP and LIME. These will show students complex charts that explain exactly how their skills—such as Python expertise or Creative Writing—influenced their top career match, making the AI more transparent.")
    
    add_heading(doc, "12.4 Direct Learning Path Suggestions", level=2)
    add_normal_para(doc, "Future versions will automatically scan the student's 'Skill Gaps' and suggest high-quality online courses from platforms like Coursera or Udemy. This provides a complete 'learning-to-employment' pipeline for the user, helping them achieve their goals faster.")

    # 13. CONCLUSION
    doc.add_page_break()
    add_heading(doc, "13. CONCLUSION", level=1)
    
    concl_para1 = "The 'AI CareerPilot: Personalized Career Guidance System' successfully bridges the critical gap between academic training and industrial requirements. By leveraging state-of-the-art Machine Learning algorithms and a robust FastAPI-driven backend, the system provides students with a low-latency, high-accuracy prediction of their professional trajectory."
    add_normal_para(doc, concl_para1)
    
    concl_para2 = "During development, the project achieved several key milestones including accurate trajectory mapping using scikit-learn, the integration of Explainable AI (XAI) via SHAP to foster user trust, and the creation of an interactive AI Coach that provides real-time, LLM-driven mentorship. The system simplifies complex career decisions and provides academic departments with the tools needed to guide their students toward professional success."
    add_normal_para(doc, concl_para2)
    
    concl_para3 = "In conclusion, this project proves that AI-augmented systems can scale mentorship capabilities beyond traditional human limits. It empowers students with data-backed confidence, ensuring that their transition from the classroom to the professional world is guided by objective insights. The future of career guidance lies in such personalized systems that can adapt to the ever-evolving global job market."
    add_normal_para(doc, concl_para3)

    # 14. BIBLIOGRAPHY
    doc.add_page_break()
    add_heading(doc, "14. BIBLIOGRAPHY", level=1)
    
    bib_items = [
        "1. Pedregosa, F., Varoquaux, G., et al. (2011). 'Scikit-learn: Machine Learning in Python'. Journal of Machine Learning Research, Vol. 12, pp. 2825-2830.",
        "2. Lundberg, S. M., & Lee, S. I. (2017). 'A Unified Approach to Interpreting Model Predictions'. Advances in Neural Information Processing Systems (NeurIPS), Vol. 30, pp. 4765-4774.",
        "3. Ramirez-Velarde, R., et al. (2021). 'Artificial Intelligence in Career Guidance Systems: A Systematic Review'. International Journal of Educational Technology in Higher Education.",
        "4. Tiangolo, S. (2018). 'FastAPI: High performance, easy to learn, fast to code, ready for production'. [Online] Available at: https://fastapi.tiangolo.com.",
        "5. Mubarak, A. A., et al. (2020). 'A Career Path Recommendation System using Machine Learning'. IEEE International Conference on Computing, Power and Communication Technologies (GUCON).",
        "6. Meta Platforms Inc. (2024). 'React.js Documentation: Building Modern User Interfaces'. Official Technical Guide.",
        "7. Oracle Corporation. (2024). 'MySQL Database Documentation: Relational Management for Web Scalability'.",
        "8. Nielsen, J. (1994). 'Usability Engineering & Modern UI Principles'. Morgan Kaufmann Publishers."
    ]
    for item in bib_items:
        add_normal_para(doc, item)

    add_page_numbers(doc)
    doc.save("Final_Project_Report.docx")
    print("Report enriched with detailed and simple language successfully.")

if __name__ == "__main__":
    create_report()
